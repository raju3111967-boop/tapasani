"""
Word Document Generator for Notices and Letters
Uses python-docx to generate .docx files from HTML templates
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup


def html_to_worddocx(html_content):
    """
    Convert HTML content to Word document
    
    Args:
        html_content (str): HTML content as string
    
    Returns:
        Document: python-docx Document object
    """
    doc = Document()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Process each element
    for element in soup.find_all(True):
        tag = element.name
        
        if tag == 'h1':
            p = doc.add_heading(element.get_text(), level=1)
        elif tag == 'h2':
            p = doc.add_heading(element.get_text(), level=2)
        elif tag == 'h3':
            p = doc.add_heading(element.get_text(), level=3)
        elif tag == 'p':
            p = doc.add_paragraph(element.get_text())
        elif tag == 'strong' or tag == 'b':
            # Bold text - handled within paragraphs
            pass
        elif tag == 'table':
            # Process table
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all('td')))
                table.style = 'Table Grid'
                
                for i, row in enumerate(rows):
                    cells = row.find_all('td')
                    for j, cell in enumerate(cells):
                        table.rows[i].cells[j].text = cell.get_text()
        elif tag == 'ul':
            # Unordered list
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif tag == 'div':
            # Div - process children
            pass
        elif tag == 'br':
            doc.add_paragraph()
    
    return doc


def generate_word_from_template(template_path, context):
    """
    Generate Word document from Django template with context
    
    Args:
        template_path (str): Path to template file
        context (dict): Context variables for template
    
    Returns:
        Document: python-docx Document object
    """
    from django.template.loader import render_to_string
    
    # Render template with context
    html_content = render_to_string(template_path, context)
    
    # Convert to Word document
    doc = html_to_worddocx(html_content)
    
    return doc


def save_word_document(doc, output_path):
    """
    Save Word document to file
    
    Args:
        doc (Document): python-docx Document object
        output_path (str): Path to save the file
    """
    doc.save(output_path)
